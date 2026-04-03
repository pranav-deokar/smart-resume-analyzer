"""
Text extraction utilities for PDF and DOCX files
"""

import pdfplumber
from docx import Document


def extract_text_from_pdf(filepath):
    """
    Extract text from PDF file using pdfplumber
    Returns cleaned text string
    """
    try:
        text = ""
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # Clean up text
        text = text.strip()
        
        if not text:
            raise ValueError("No text could be extracted from PDF")
        
        return text
    
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(filepath):
    """
    Extract text from DOCX file using python-docx
    Returns cleaned text string
    """
    try:
        doc = Document(filepath)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Clean up text
        text = text.strip()
        
        if not text:
            raise ValueError("No text could be extracted from DOCX")
        
        return text
    
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")
